from docx import Document

doc = Document()
doc.add_heading('OWASP Top 10 Security Scan Report', 0)

doc.add_heading('Semgrep Report (Static Scan)', level=1)
with open('reports/semgrep_report.txt') as f:
    for line in f:
        doc.add_paragraph(line.strip())
 #this is going to store the file        
doc.add_paragraph('For detailed issues, please see the attached zap_report.html.')

doc.save('reports/security_scan_report.docx')

